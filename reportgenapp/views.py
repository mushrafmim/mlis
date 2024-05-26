import datetime
import docx
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status


@api_view(['POST'])
def generate_report(request):
    # Generate the report
    patientData = request.data.get('patientData')

    replacements = {
        "{name}": patientData['name'],
        "{age}": patientData['age'],
        "{referredBy}": patientData['referredBy'],
        "{date}": datetime.datetime.now().strftime("%Y-%m-%d"),
        "{labNo}": "23044",
        "{result}": "125",
        "{remarks}": "NORMAL",
    }

    doc = docx.Document("reportgenapp/templates/fbs.docx")

    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)

    doc.save("reportgenapp/reports/fbs_report.docx")
    return Response({'message': 'Report generated successfully'}, status=status.HTTP_200_OK)
