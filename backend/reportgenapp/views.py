import os
import pythoncom
import datetime
import docx2pdf
import docx
from django.http import FileResponse
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from rest_framework.views import APIView

from .models import ReportRequest, Report, ReportFormat
from reportgenapp.serializers.serializers import (AddToReportQueueRequestSerializer, GetReportQueueResponseSerializer,
                                                  GetReportFormatsResponseSerializer, GetReportRequestItemSerializer)
from reportgenapp.serializers.report_format import ReportFormatUpdateValuesSerializer
from .models import REPORTSTATUS


class ReportFormatAPIView(APIView):

    def get(self, request):
        # Get the report formats
        report_formats = ReportFormat.objects.all()

        serialized = GetReportFormatsResponseSerializer(report_formats, many=True)
        return Response(serialized.data, status=status.HTTP_200_OK)

    def put(self, request, id):
        # Update the report format
        report_format = ReportFormat.objects.filter(id=id).first()

        if not report_format:
            return Response({'error': 'Report format not found'}, status=status.HTTP_404_NOT_FOUND)

        serialized = ReportFormatUpdateValuesSerializer(report_format, data=request.data, partial=True)

        if not serialized.is_valid():
            return Response(serialized.errors, status=status.HTTP_400_BAD_REQUEST)

        serialized.save()
        return Response(serialized.data, status=status.HTTP_200_OK)


class ReportQueueAPIView(APIView):

    def get(self, request):
        # Get the report queue
        report_queue = (Report.objects.filter(status__in=[REPORTSTATUS.REPORT_PENDING])
                        .order_by('-created_at').all())

        serialized = GetReportQueueResponseSerializer(report_queue, many=True)
        return Response(serialized.data, status=status.HTTP_200_OK)

    def post(self, request):
        # Add the report to the queue
        # serialize the patient data
        serialized = AddToReportQueueRequestSerializer(data=request.data)

        if not serialized.is_valid():
            return Response(serialized.errors, status=status.HTTP_400_BAD_REQUEST)

        patient_id = serialized.data['patient']
        name = serialized.data['name']
        age = serialized.data['age']
        gender = serialized.data['gender']
        reports = serialized.data['reports']

        # loop through the report formats and save the report to the queue
        request = ReportRequest.objects.create(
            patient_id=patient_id,
            name=name,
            age=age,
            gender=gender
        )

        for report in reports:
            # Get the report format
            report_format = ReportFormat.objects.filter(id=report['report_format']).first()

            delivery_date = datetime.datetime.now(datetime.timezone.utc) + datetime.timedelta(
                minutes=report_format.generation_time)

            Report.objects.create(
                report_format_id=report['report_format'],
                report_request=request,
                referred_by=report['referred_by'],
                delivery_date=delivery_date,
                has_paid=report['has_paid']
            )

        return Response(status=status.HTTP_201_CREATED)


class ReportQueueItemAPIView(APIView):

    def get(self, request, id):
        # Get the report queue item
        report_queue_item = ReportRequest.objects.filter(id=id).first()

        if not report_queue_item:
            return Response({'error': 'Report queue item not found'}, status=status.HTTP_404_NOT_FOUND)

        serialized = GetReportRequestItemSerializer(report_queue_item)
        return Response(serialized.data, status=status.HTTP_200_OK)

    def post(self, request, id, report_format):
        # Update the report queue item
        print(id)
        report_queue_item = ReportRequest.objects.filter(id=id).first()

        if not report_queue_item:
            return Response({'error': 'Report queue item not found'}, status=status.HTTP_404_NOT_FOUND)

        report = Report.objects.filter(report_request_id=id).first()

        if not report:
            return Response({'error': 'Report not found'}, status=status.HTTP_404_NOT_FOUND)

        report.status = REPORTSTATUS.REPORT_GENERATED
        report.values = request.data['values']
        report.save()

        return Response(status=status.HTTP_200_OK)


@api_view(['POST'])
def mark_as_delivered(request, id):
    report = Report.objects.filter(id=id).first()

    if not report:
        return Response({'error': 'Report not found'}, status=status.HTTP_404_NOT_FOUND)

    report.status = REPORTSTATUS.REPORT_DELIVERED
    report.save()

    return Response(status=status.HTTP_200_OK)


@api_view(['POST'])
def generate_report(request, report_id):
    try:
        # Co Initialization
        pythoncom.CoInitialize()
        # Generate the report

        report = Report.objects.filter(id=report_id).first()

        if not report:
            return Response({'error': 'Report not found'}, status=status.HTTP_404_NOT_FOUND)

        report_format = report.report_format

        replacements = report_format.replacements

        # serial number of the report
        now = datetime.datetime.now()
        serial_no = now.strftime('%y%m%d%H%M%S') + f"{now.microsecond // 1000:03d}"

        # Adding date to the replacements
        replacements['{date}'] = datetime.datetime.now().strftime('%d-%m-%Y')
        replacements['{labNo}'] = serial_no
        replacements['{referredBy}'] = report.referred_by

        # Loop through the replacements and replace the placeholders in the template
        # with the actual values

        # replacing the report data
        report_dict = report.report_request.__dict__
        for key, value in replacements.items():
            if value in report_dict:
                replacements[key] = report_dict[value]

        print(report_dict)

        # replacing the report values
        for key, value in replacements.items():
            if value in report.values:
                replacements[key] = str(report.values[value])

        # pic the file path of the corresponding template
        template_path = report_format.template_path

        # Load the template
        doc = docx.Document(template_path)

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)

        os.makedirs('reportgenapp/generated', exist_ok=True)

        # Generate a random file name
        file_name = f"reportgenapp/generated/{serial_no}.docx"

        doc.save(file_name)

        docx2pdf.convert(file_name)

        # Delete the docx file
        os.remove(file_name)

        pdf_path = file_name.replace('.docx', '.pdf')

        return FileResponse(open(pdf_path, 'rb'), content_type='application/pdf', as_attachment=True,
                            filename='document.pdf')

    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
